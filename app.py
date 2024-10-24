import streamlit as st
from docx import Document
from docx.shared import Pt
import io
import requests

# Configuración de la página
st.set_page_config(
    page_title="📄 Enriquecedor de Documentos DOCX",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("📄 Enriquecedor de Documentos DOCX")
st.write("""
    Sube tu documento en formato DOCX, y esta aplicación agregará subtítulos y enriquecerá su contenido utilizando las APIs de Serper y OpenRouter.
""")

# Subida de archivo
uploaded_file = st.file_uploader("Sube tu documento DOCX", type=["docx"])

if uploaded_file is not None:
    try:
        # Leer el documento DOCX
        doc = Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        original_text = '\n'.join(full_text)

        st.subheader("Contenido Original del Documento")
        st.text_area("", original_text, height=300)

        # Función para realizar búsquedas con Serper API
        def search_with_serper(query):
            serper_api_key = st.secrets["SERPER_API_KEY"]
            url = "https://google.serper.dev/search"
            headers = {
                "X-API-KEY": serper_api_key,
                "Content-Type": "application/json"
            }
            data = {
                "q": query
            }
            response = requests.post(url, headers=headers, json=data)
            if response.status_code == 200:
                return response.json()
            else:
                st.error(f"Error en la API de Serper: {response.status_code}")
                st.error(response.text)
                return None

        # Extraer temas clave del documento para realizar búsquedas relevantes
        def extract_topics(text, max_topics=5):
            prompt = (
                "Extrae los temas o palabras clave más relevantes del siguiente texto, separados por comas:\n\n"
                f"{text}"
            )
            api_key = st.secrets["OPENROUTER_API_KEY"]
            url = "https://openrouter.ai/api/v1/chat/completions"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            data = {
                "model": "openai/gpt-4o-mini",
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            }
            response = requests.post(url, headers=headers, json=data)
            if response.status_code == 200:
                response_json = response.json()
                topics = response_json['choices'][0]['message']['content']
                # Limitar el número de temas
                topics_list = [topic.strip() for topic in topics.split(',')][:max_topics]
                return topics_list
            else:
                st.error(f"Error en la API de OpenRouter al extraer temas: {response.status_code}")
                st.error(response.text)
                return []

        with st.spinner("Extrayendo temas clave del documento..."):
            topics = extract_topics(original_text)
        
        if topics:
            st.subheader("Temas Clave Extraídos")
            st.write(", ".join(topics))

            # Realizar búsquedas en Serper para cada tema
            search_results = {}
            for topic in topics:
                with st.spinner(f"Realizando búsqueda para: {topic}"):
                    result = search_with_serper(topic)
                    if result:
                        search_results[topic] = result

            # Preparar la información obtenida de las búsquedas para el enriquecimiento
            enrichment_data = ""
            for topic, data in search_results.items():
                enrichment_data += f"\n\n### {topic}\n"
                # Agregar snippets de las búsquedas
                if 'organic' in data:
                    for item in data['organic'][:3]:  # Limitar a los primeros 3 resultados
                        enrichment_data += f"- {item.get('title')}: {item.get('snippet')}\n"
                else:
                    enrichment_data += "No se encontraron resultados relevantes.\n"

            # Preparar el prompt para la API de OpenRouter con los datos de enriquecimiento
            prompt = (
                "Agrega subtítulos y enriquece el siguiente documento con datos adicionales donde sea posible. Utiliza la información adicional proporcionada a continuación:\n\n"
                f"{original_text}\n\n"
                f"Datos Adicionales:{enrichment_data}"
            )

            # Función para obtener el texto enriquecido usando OpenRouter
            def get_enriched_text(prompt):
                api_key = st.secrets["OPENROUTER_API_KEY"]
                url = "https://openrouter.ai/api/v1/chat/completions"
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}"
                }
                data = {
                    "model": "openai/gpt-4o-mini",
                    "messages": [
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ]
                }
                response = requests.post(url, headers=headers, json=data)
                if response.status_code == 200:
                    response_json = response.json()
                    return response_json['choices'][0]['message']['content']
                else:
                    st.error(f"Error en la API de OpenRouter: {response.status_code}")
                    st.error(response.text)
                    return None

            with st.spinner("Enriqueciendo el documento..."):
                enriched_text = get_enriched_text(prompt)

            if enriched_text:
                st.subheader("Contenido Enriquecido del Documento")
                st.text_area("", enriched_text, height=300)

                # Crear un nuevo documento DOCX con el contenido enriquecido
                new_doc = Document()
                for line in enriched_text.split('\n'):
                    if line.strip().startswith("### "):
                        # Subtítulos de nivel 2
                        new_doc.add_heading(line.replace("### ", ""), level=2)
                    elif line.strip().startswith("## "):
                        # Subtítulos de nivel 1
                        new_doc.add_heading(line.replace("## ", ""), level=1)
                    else:
                        new_doc.add_paragraph(line)

                # Guardar el nuevo documento en un objeto BytesIO
                byte_io = io.BytesIO()
                new_doc.save(byte_io)
                byte_io.seek(0)

                st.success("¡Documento enriquecido creado con éxito!")
                st.download_button(
                    label="📥 Descargar Documento Enriquecido",
                    data=byte_io,
                    file_name="documento_enriquecido.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("No se pudieron extraer temas clave del documento para realizar búsquedas.")
    except Exception as e:
        st.error(f"Ocurrió un error al procesar el documento: {e}")
